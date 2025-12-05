"""
AI Sales Copy & Ad Content Agent
================================
A complete AI system that generates multi-platform marketing content including:
- Google Ad headlines & descriptions
- Facebook/Instagram ad copies
- Captions + Hashtags
- SEO titles & meta descriptions
- CTA suggestions
- Keyword extraction
- Landing page content

Author: AI Sales Agent
Version: 3.0.0 (Premium UI Edition - Groq FREE)
"""

import streamlit as st
from groq import Groq
import os
import json
import sqlite3
import re
import io
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import pandas as pd
import nltk
from collections import Counter

# Load environment variables
load_dotenv()

# =============================================================================
# PREMIUM CSS STYLES
# =============================================================================

PREMIUM_CSS = """
<style>
/* Import Google Fonts */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=Poppins:wght@400;500;600;700;800&display=swap');

/* Root Variables */
:root {
    --primary-gradient: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    --secondary-gradient: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
    --success-gradient: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
    --dark-gradient: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
    --card-shadow: 0 10px 40px rgba(0,0,0,0.1);
    --hover-shadow: 0 20px 60px rgba(0,0,0,0.15);
}

/* Hide Streamlit Branding */
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}

/* Main Container */
.main .block-container {
    padding: 2rem 3rem;
    max-width: 1400px;
}

/* Premium Header */
.premium-header {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
    padding: 3rem 2rem;
    border-radius: 24px;
    margin-bottom: 2rem;
    text-align: center;
    box-shadow: 0 20px 60px rgba(102, 126, 234, 0.4);
    position: relative;
    overflow: hidden;
}

.premium-header::before {
    content: '';
    position: absolute;
    top: -50%;
    left: -50%;
    width: 200%;
    height: 200%;
    background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 60%);
    animation: pulse 4s ease-in-out infinite;
}

@keyframes pulse {
    0%, 100% { transform: scale(1); opacity: 0.5; }
    50% { transform: scale(1.1); opacity: 0.8; }
}

.premium-header h1 {
    color: white;
    font-family: 'Poppins', sans-serif;
    font-size: 3rem;
    font-weight: 800;
    margin: 0;
    text-shadow: 2px 4px 20px rgba(0,0,0,0.3);
    position: relative;
    z-index: 1;
}

.premium-header p {
    color: rgba(255,255,255,0.95);
    font-size: 1.3rem;
    margin-top: 1rem;
    font-weight: 400;
    position: relative;
    z-index: 1;
}

.premium-badge {
    display: inline-block;
    background: rgba(255,255,255,0.25);
    backdrop-filter: blur(10px);
    padding: 0.5rem 1.5rem;
    border-radius: 50px;
    color: white;
    font-weight: 600;
    font-size: 0.9rem;
    margin-top: 1rem;
    border: 1px solid rgba(255,255,255,0.3);
}

/* Feature Cards */
.feature-card {
    background: white;
    border-radius: 20px;
    padding: 2rem;
    box-shadow: 0 10px 40px rgba(0,0,0,0.08);
    transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
    border: 1px solid rgba(0,0,0,0.05);
    height: 100%;
}

.feature-card:hover {
    transform: translateY(-10px);
    box-shadow: 0 25px 60px rgba(102, 126, 234, 0.2);
}

.feature-icon {
    font-size: 3rem;
    margin-bottom: 1rem;
}

.feature-title {
    font-family: 'Poppins', sans-serif;
    font-size: 1.4rem;
    font-weight: 700;
    color: #1a1a2e;
    margin-bottom: 0.8rem;
}

.feature-desc {
    color: #666;
    font-size: 0.95rem;
    line-height: 1.6;
}

/* Stats Cards */
.stats-card {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    border-radius: 16px;
    padding: 1.5rem;
    text-align: center;
    color: white;
    box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
}

.stats-number {
    font-size: 2.5rem;
    font-weight: 800;
    font-family: 'Poppins', sans-serif;
}

.stats-label {
    font-size: 0.9rem;
    opacity: 0.9;
    margin-top: 0.3rem;
}

/* Premium Buttons */
.stButton>button {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    color: white;
    border: none;
    border-radius: 12px;
    padding: 0.8rem 2rem;
    font-weight: 600;
    font-size: 1rem;
    transition: all 0.3s ease;
    box-shadow: 0 8px 25px rgba(102, 126, 234, 0.35);
}

.stButton>button:hover {
    transform: translateY(-3px);
    box-shadow: 0 15px 35px rgba(102, 126, 234, 0.45);
}

.stButton>button:active {
    transform: translateY(-1px);
}

/* Input Fields */
.stTextInput>div>div>input,
.stTextArea>div>div>textarea {
    border-radius: 12px;
    border: 2px solid #e0e0e0;
    padding: 0.8rem 1rem;
    font-size: 1rem;
    transition: all 0.3s ease;
}

.stTextInput>div>div>input:focus,
.stTextArea>div>div>textarea:focus {
    border-color: #667eea;
    box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.15);
}

/* Select Box */
.stSelectbox>div>div {
    border-radius: 12px;
}

/* Expander */
div[data-testid="stExpander"] {
    background: white;
    border-radius: 16px;
    border: none;
    box-shadow: 0 5px 20px rgba(0,0,0,0.08);
    margin-bottom: 1rem;
    overflow: hidden;
}

div[data-testid="stExpander"] div[role="button"] {
    padding: 1rem 1.5rem;
}

div[data-testid="stExpander"] div[role="button"] p {
    font-size: 1.1rem;
    font-weight: 600;
    color: #1a1a2e !important;
}

/* Force dark text in content areas */
div[data-testid="stExpander"] .stMarkdown p,
div[data-testid="stExpander"] .stMarkdown span,
div[data-testid="stExpander"] .stMarkdown li {
    color: #1a1a2e !important;
}

div[data-testid="stExpander"] .stMarkdown h1,
div[data-testid="stExpander"] .stMarkdown h2,
div[data-testid="stExpander"] .stMarkdown h3,
div[data-testid="stExpander"] .stMarkdown h4 {
    color: #667eea !important;
}

div[data-testid="stExpander"] .stMarkdown strong {
    color: #667eea !important;
}

/* Main content area text */
.main .block-container p,
.main .block-container span,
.main .block-container li {
    color: #1a1a2e;
}

/* Ensure readability in all cards */
.feature-card p,
.feature-card span,
.step-card p,
.result-card p {
    color: #333 !important;
}

/* Content Results Card */
.result-card {
    background: linear-gradient(145deg, #ffffff 0%, #f8f9ff 100%);
    border-radius: 20px;
    padding: 2rem;
    margin: 1rem 0;
    box-shadow: 0 10px 40px rgba(0,0,0,0.06);
    border-left: 5px solid #667eea;
}

.result-title {
    font-family: 'Poppins', sans-serif;
    font-size: 1.3rem;
    font-weight: 700;
    color: #667eea;
    margin-bottom: 1rem;
    display: flex;
    align-items: center;
    gap: 0.5rem;
}

/* Sidebar Styling */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1a1a2e 0%, #16213e 100%);
}

section[data-testid="stSidebar"] .stMarkdown {
    color: white;
}

section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: white !important;
}

section[data-testid="stSidebar"] .stRadio label {
    color: rgba(255,255,255,0.9) !important;
}

section[data-testid="stSidebar"] .stTextInput label {
    color: rgba(255,255,255,0.9) !important;
}

/* Download Buttons */
.stDownloadButton>button {
    background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%);
    color: white;
    border: none;
    border-radius: 12px;
    padding: 0.7rem 1.5rem;
    font-weight: 600;
    box-shadow: 0 8px 25px rgba(17, 153, 142, 0.3);
}

.stDownloadButton>button:hover {
    transform: translateY(-2px);
    box-shadow: 0 12px 30px rgba(17, 153, 142, 0.4);
}

/* Metrics */
div[data-testid="metric-container"] {
    background: white;
    border-radius: 16px;
    padding: 1.5rem;
    box-shadow: 0 8px 30px rgba(0,0,0,0.08);
}

/* Success/Info/Warning boxes */
.stSuccess, .stInfo, .stWarning {
    border-radius: 12px;
    padding: 1rem 1.5rem;
}

/* Animations */
@keyframes fadeInUp {
    from {
        opacity: 0;
        transform: translateY(30px);
    }
    to {
        opacity: 1;
        transform: translateY(0);
    }
}

.animate-in {
    animation: fadeInUp 0.6s ease-out forwards;
}

/* Steps */
.step-card {
    background: white;
    border-radius: 20px;
    padding: 2rem;
    text-align: center;
    box-shadow: 0 10px 40px rgba(0,0,0,0.08);
    transition: all 0.3s ease;
    height: 100%;
}

.step-card:hover {
    transform: translateY(-5px);
    box-shadow: 0 20px 50px rgba(0,0,0,0.12);
}

.step-number {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    color: white;
    width: 50px;
    height: 50px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 1.5rem;
    font-weight: 700;
    margin: 0 auto 1rem;
}

.step-title {
    font-family: 'Poppins', sans-serif;
    font-size: 1.2rem;
    font-weight: 700;
    color: #1a1a2e;
    margin-bottom: 0.5rem;
}

/* Glow effect */
.glow {
    animation: glow 2s ease-in-out infinite alternate;
}

@keyframes glow {
    from {
        box-shadow: 0 0 20px rgba(102, 126, 234, 0.4);
    }
    to {
        box-shadow: 0 0 40px rgba(102, 126, 234, 0.8);
    }
}

/* Floating animation */
.float {
    animation: float 3s ease-in-out infinite;
}

@keyframes float {
    0%, 100% { transform: translateY(0); }
    50% { transform: translateY(-10px); }
}

/* Gradient text */
.gradient-text {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
}

/* Platform tags */
.platform-tag {
    display: inline-block;
    padding: 0.4rem 1rem;
    border-radius: 50px;
    font-size: 0.85rem;
    font-weight: 600;
    margin: 0.3rem;
}

.tag-google { background: linear-gradient(135deg, #4285f4, #34a853); color: white; }
.tag-facebook { background: linear-gradient(135deg, #1877f2, #3b5998); color: white; }
.tag-instagram { background: linear-gradient(135deg, #f09433, #e6683c, #dc2743, #cc2366, #bc1888); color: white; }
.tag-seo { background: linear-gradient(135deg, #11998e, #38ef7d); color: white; }
.tag-landing { background: linear-gradient(135deg, #667eea, #764ba2); color: white; }

/* Scrollbar */
::-webkit-scrollbar {
    width: 8px;
    height: 8px;
}

::-webkit-scrollbar-track {
    background: #f1f1f1;
    border-radius: 10px;
}

::-webkit-scrollbar-thumb {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    border-radius: 10px;
}

/* Mobile Responsive */
@media (max-width: 768px) {
    .premium-header h1 {
        font-size: 2rem;
    }
    .premium-header p {
        font-size: 1rem;
    }
    .main .block-container {
        padding: 1rem;
    }
}
</style>
"""

# =============================================================================
# DATABASE SETUP
# =============================================================================

def init_database():
    """Initialize SQLite database with required tables"""
    conn = sqlite3.connect('sales_content.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            email TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS content_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            business_name TEXT,
            business_type TEXT,
            product_service TEXT,
            target_audience TEXT,
            offer TEXT,
            tone TEXT,
            platform TEXT,
            headlines TEXT,
            descriptions TEXT,
            hashtags TEXT,
            keywords TEXT,
            cta TEXT,
            seo_title TEXT,
            meta_description TEXT,
            landing_page_content TEXT,
            full_response TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    conn.commit()
    conn.close()

def save_to_history(user_id, inputs, outputs):
    """Save generated content to database history"""
    conn = sqlite3.connect('sales_content.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO content_history 
        (user_id, business_name, business_type, product_service, target_audience, 
         offer, tone, platform, headlines, descriptions, hashtags, keywords, 
         cta, seo_title, meta_description, landing_page_content, full_response)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        user_id,
        inputs.get('business_name', ''),
        inputs.get('business_type', ''),
        inputs.get('product_service', ''),
        inputs.get('target_audience', ''),
        inputs.get('offer', ''),
        inputs.get('tone', ''),
        inputs.get('platform', ''),
        outputs.get('headlines', ''),
        outputs.get('descriptions', ''),
        outputs.get('hashtags', ''),
        outputs.get('keywords', ''),
        outputs.get('cta', ''),
        outputs.get('seo_title', ''),
        outputs.get('meta_description', ''),
        outputs.get('landing_page_content', ''),
        json.dumps(outputs)
    ))
    
    conn.commit()
    conn.close()

def get_user_history(user_id, limit=50):
    """Retrieve user's content generation history"""
    conn = sqlite3.connect('sales_content.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT * FROM content_history 
        WHERE user_id = ? 
        ORDER BY created_at DESC 
        LIMIT ?
    ''', (user_id, limit))
    
    columns = [description[0] for description in cursor.description]
    rows = cursor.fetchall()
    conn.close()
    
    return [dict(zip(columns, row)) for row in rows]

# =============================================================================
# NLP KEYWORD EXTRACTION ENGINE
# =============================================================================

def download_nltk_data():
    """Download required NLTK data"""
    nltk_packages = [
        ('tokenizers/punkt', 'punkt'),
        ('tokenizers/punkt_tab', 'punkt_tab'),
        ('corpora/stopwords', 'stopwords'),
        ('taggers/averaged_perceptron_tagger', 'averaged_perceptron_tagger'),
        ('taggers/averaged_perceptron_tagger_eng', 'averaged_perceptron_tagger_eng'),
    ]
    
    for path, package in nltk_packages:
        try:
            nltk.data.find(path)
        except LookupError:
            nltk.download(package, quiet=True)

def extract_keywords_nlp(text, num_keywords=15):
    """Extract keywords from text using NLTK"""
    download_nltk_data()
    
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    from nltk import pos_tag
    
    tokens = word_tokenize(text.lower())
    stop_words = set(stopwords.words('english'))
    
    marketing_stopwords = {
        'will', 'can', 'get', 'make', 'use', 'new', 'one', 'also', 
        'like', 'just', 'know', 'take', 'come', 'see', 'want', 'look',
        'give', 'think', 'good', 'best', 'way', 'need', 'feel', 'try'
    }
    stop_words.update(marketing_stopwords)
    
    filtered_tokens = [
        token for token in tokens 
        if token.isalnum() and token not in stop_words and len(token) > 2
    ]
    
    pos_tags = pos_tag(filtered_tokens)
    important_tags = {'NN', 'NNS', 'NNP', 'NNPS', 'VB', 'VBG', 'JJ', 'JJR', 'JJS'}
    
    important_words = [
        word for word, tag in pos_tags 
        if tag in important_tags
    ]
    
    word_freq = Counter(important_words)
    keywords = [word for word, count in word_freq.most_common(num_keywords)]
    
    return keywords

def generate_hashtags(keywords, platform='instagram'):
    """Generate platform-appropriate hashtags from keywords"""
    hashtags = []
    
    for keyword in keywords[:10]:
        tag = keyword.replace(' ', '').replace('-', '').lower()
        if tag:
            hashtags.append(f"#{tag}")
    
    if platform.lower() in ['instagram', 'facebook']:
        common_tags = ['#marketing', '#business', '#entrepreneur', '#success', '#growth']
        hashtags.extend(common_tags[:3])
    
    return list(set(hashtags))[:15]

# =============================================================================
# PROMPT TEMPLATES ENGINE (ENHANCED FOR BETTER HEADLINES)
# =============================================================================

class PromptTemplates:
    """Enhanced prompt templates for high-converting content"""
    
    @staticmethod
    def get_tone_modifier(tone):
        """Return tone-specific writing instructions"""
        tone_modifiers = {
            'Professional': "Use formal, business-appropriate language. Be authoritative and trustworthy. Focus on value propositions and credibility.",
            'Emotional': "Connect emotionally with the reader. Use storytelling elements. Appeal to feelings, desires, and aspirations. Use power words that trigger emotions.",
            'Exciting': "Use energetic, dynamic language. Create enthusiasm and anticipation. Use action words and create FOMO (fear of missing out).",
            'Urgent': "Create a sense of urgency and scarcity. Use time-sensitive language. Emphasize limited availability or time-bound offers. Use words like NOW, TODAY, LIMITED, HURRY.",
            'Friendly': "Use warm, conversational tone. Be approachable and relatable. Write as if talking to a friend. Use casual language.",
            'Luxury': "Use sophisticated, premium language. Emphasize exclusivity and quality. Appeal to aspirational desires. Use words like EXCLUSIVE, PREMIUM, ELITE."
        }
        return tone_modifiers.get(tone, tone_modifiers['Professional'])
    
    @staticmethod
    def google_ads_prompt(inputs):
        """Generate Google Ads content prompt with high-converting headlines"""
        return f"""
You are a world-class Google Ads copywriter who has generated millions in revenue. Create IRRESISTIBLE, HIGH-CONVERTING Google Ads content.

BUSINESS DETAILS:
- Business Name: {inputs['business_name']}
- Business Type: {inputs['business_type']}
- Product/Service: {inputs['product_service']}
- Target Audience: {inputs['target_audience']}
- Offer: {inputs['offer']}
- Tone: {inputs['tone']}

TONE INSTRUCTIONS: {PromptTemplates.get_tone_modifier(inputs['tone'])}

HEADLINE POWER FORMULAS (Use these patterns):
1. [Number] + [Benefit] + [Timeframe] → "Get 3X Sales in 30 Days"
2. [Action Verb] + [Desire] + [Differentiator] → "Unlock Premium Results Fast"
3. [Question] that triggers curiosity → "Want 10X More Leads?"
4. [Pain Point] + [Solution] → "Tired of Low Sales? Fix It Now"
5. [Social Proof] + [Result] → "Join 10K+ Happy Customers"
6. [Urgency] + [Benefit] → "Limited: 50% Off Today Only"

POWER WORDS TO USE: Free, New, Proven, Guaranteed, Instant, Exclusive, Secret, Easy, Fast, Save, Discover, Unlock, Transform, Boost, Skyrocket

STRICT CHARACTER LIMITS:
- Headlines: Maximum 30 characters each (including spaces)
- Descriptions: Maximum 90 characters each (including spaces)

Generate the following in JSON format:
{{
    "headlines": [
        // 15 IRRESISTIBLE headlines using power formulas, each MUST be 30 characters or less
        // Mix different formulas for variety
        // Include numbers, power words, and emotional triggers
    ],
    "descriptions": [
        // 5 compelling descriptions with clear value proposition and CTA, max 90 chars each
    ],
    "display_urls": [
        // 3 keyword-rich display URL paths
    ],
    "keywords": [
        // 15 high-intent search keywords
    ],
    "negative_keywords": [
        // 5 negative keywords to exclude
    ],
    "cta_suggestions": [
        // 5 action-oriented CTAs
    ]
}}

IMPORTANT: 
- Every headline must trigger curiosity or desire
- Use numbers and specifics (3X, 50%, 24hrs)
- Include at least 3 headlines with urgency
- Count characters carefully - headlines over 30 chars will be rejected
- Return ONLY valid JSON
"""

    @staticmethod
    def facebook_instagram_prompt(inputs):
        """Generate Facebook/Instagram ad content prompt"""
        return f"""
You are a viral social media marketing expert with 10M+ followers experience. Create SCROLL-STOPPING content.

BUSINESS DETAILS:
- Business Name: {inputs['business_name']}
- Business Type: {inputs['business_type']}
- Product/Service: {inputs['product_service']}
- Target Audience: {inputs['target_audience']}
- Offer: {inputs['offer']}
- Tone: {inputs['tone']}

TONE INSTRUCTIONS: {PromptTemplates.get_tone_modifier(inputs['tone'])}

VIRAL HOOKS FORMULAS:
1. "Stop scrolling if you..." 
2. "POV: You just discovered..."
3. "This is your sign to..."
4. "Nobody talks about this but..."
5. "The secret to [desire] is..."
6. "Here's why [common belief] is wrong..."

Generate the following in JSON format:
{{
    "facebook_ad": {{
        "primary_text": [
            // 3 SCROLL-STOPPING primary texts with hooks (125-500 chars)
            // Start with a hook that stops the scroll
            // Include social proof and urgency
        ],
        "headlines": [
            // 5 curiosity-driven headlines (max 40 characters)
        ],
        "descriptions": [
            // 3 benefit-focused link descriptions (max 30 characters)
        ],
        "cta_button": [
            "Shop Now", "Learn More", "Sign Up", "Get Offer", "Book Now"
        ]
    }},
    "instagram_ad": {{
        "captions": [
            // 3 engaging captions with emojis, hooks, and story elements
            // Use line breaks for readability
            // End with strong CTA
        ],
        "story_text": [
            // 3 punchy story overlay texts (max 100 characters)
        ],
        "hashtags": [
            // 25 strategic hashtags (mix of popular, niche, and branded)
        ],
        "bio_link_cta": [
            // 3 compelling "Link in bio" CTAs
        ],
        "reels_hooks": [
            // 5 viral reel opening hooks
        ]
    }},
    "carousel_hooks": [
        // 5 carousel slide headline hooks that make people swipe
    ],
    "engagement_questions": [
        // 3 questions to boost comments and engagement
    ]
}}

IMPORTANT:
- First line MUST stop the scroll
- Use emojis strategically (not excessively)
- Create FOMO and urgency
- Return ONLY valid JSON
"""

    @staticmethod
    def seo_content_prompt(inputs):
        """Generate SEO-optimized content prompt"""
        return f"""
You are an SEO expert who has ranked 1000+ pages on Google's first page.

BUSINESS DETAILS:
- Business Name: {inputs['business_name']}
- Business Type: {inputs['business_type']}
- Product/Service: {inputs['product_service']}
- Target Audience: {inputs['target_audience']}
- Offer: {inputs['offer']}
- Tone: {inputs['tone']}

TONE INSTRUCTIONS: {PromptTemplates.get_tone_modifier(inputs['tone'])}

SEO TITLE FORMULAS:
1. [Primary Keyword] - [Benefit] | [Brand]
2. [Number] Best [Keyword] for [Audience] in [Year]
3. How to [Achieve Result] with [Solution] [Year]
4. [Keyword]: The Ultimate Guide to [Benefit]

Generate the following in JSON format:
{{
    "seo_titles": [
        // 5 click-worthy SEO titles (50-60 characters) using formulas above
    ],
    "meta_descriptions": [
        // 5 compelling meta descriptions with CTA (150-160 characters)
    ],
    "h1_headings": [
        // 3 powerful H1 headings with primary keyword
    ],
    "h2_subheadings": [
        // 5 engaging H2 subheadings
    ],
    "primary_keywords": [
        // 5 high-volume primary keywords
    ],
    "secondary_keywords": [
        // 10 LSI/secondary keywords
    ],
    "long_tail_keywords": [
        // 10 long-tail keyword phrases with buyer intent
    ],
    "url_slugs": [
        // 3 SEO-friendly URL slugs
    ],
    "image_alt_texts": [
        // 5 descriptive image alt texts
    ],
    "schema_suggestions": {{
        "type": "suggested schema type",
        "key_properties": ["list of key schema properties"]
    }}
}}

Return ONLY valid JSON.
"""

    @staticmethod
    def landing_page_prompt(inputs):
        """Generate landing page content prompt"""
        return f"""
You are a conversion rate optimization expert with a track record of 40%+ conversion rates.

BUSINESS DETAILS:
- Business Name: {inputs['business_name']}
- Business Type: {inputs['business_type']}
- Product/Service: {inputs['product_service']}
- Target Audience: {inputs['target_audience']}
- Offer: {inputs['offer']}
- Tone: {inputs['tone']}

TONE INSTRUCTIONS: {PromptTemplates.get_tone_modifier(inputs['tone'])}

HEADLINE FORMULAS FOR HIGH CONVERSION:
1. "[Result] Without [Pain Point]"
2. "The [Adjective] Way to [Achieve Desire]"
3. "Get [Specific Result] in [Timeframe] or [Guarantee]"
4. "Finally, [Solution] That [Unique Benefit]"

Generate the following in JSON format:
{{
    "hero_section": {{
        "headline": "Powerful headline using formula above (max 10 words)",
        "subheadline": "Supporting text that expands the promise (max 20 words)",
        "cta_button_text": "Action-oriented CTA (e.g., 'Start Free Trial')",
        "cta_supporting_text": "Risk reducer (e.g., 'No credit card required • Cancel anytime')"
    }},
    "value_propositions": [
        {{
            "title": "Benefit-focused title",
            "description": "2-3 sentence description with specific outcomes",
            "icon_suggestion": "relevant icon name"
        }}
        // 4 total value propositions
    ],
    "features_benefits": [
        {{
            "feature": "Feature name",
            "benefit": "What it means for the user (outcome-focused)"
        }}
        // 6 feature-benefit pairs
    ],
    "social_proof": {{
        "testimonial_prompts": [
            // 3 testimonial templates
        ],
        "stats_suggestions": [
            // 3 impressive stats to showcase
        ],
        "trust_badges": [
            // 5 trust elements
        ]
    }},
    "faq_questions": [
        {{
            "question": "Common objection as question",
            "answer": "Objection-handling answer"
        }}
        // 5 FAQ items that handle objections
    ],
    "urgency_elements": [
        // 3 urgency/scarcity elements
    ],
    "final_cta": {{
        "headline": "Final push headline",
        "cta_text": "Strong final CTA",
        "guarantee": "Risk reversal guarantee"
    }}
}}

Return ONLY valid JSON.
"""

    @staticmethod
    def multi_platform_prompt(inputs):
        """Generate content for all platforms at once"""
        return f"""
You are a multi-channel marketing genius who has scaled brands from 0 to millions.

BUSINESS DETAILS:
- Business Name: {inputs['business_name']}
- Business Type: {inputs['business_type']}
- Product/Service: {inputs['product_service']}
- Target Audience: {inputs['target_audience']}
- Offer: {inputs['offer']}
- Tone: {inputs['tone']}

TONE INSTRUCTIONS: {PromptTemplates.get_tone_modifier(inputs['tone'])}

Generate IRRESISTIBLE, HIGH-CONVERTING content for ALL platforms in JSON format:
{{
    "google_ads": {{
        "headlines": [
            // 15 POWER headlines using formulas: numbers, benefits, urgency
            // Max 30 chars each - COUNT CAREFULLY
        ],
        "descriptions": [
            // 5 compelling descriptions, max 90 chars each
        ],
        "keywords": ["15 high-intent keywords"]
    }},
    "facebook": {{
        "primary_texts": [
            // 3 scroll-stopping ad texts with hooks
        ],
        "headlines": ["5 curiosity-driven headlines, max 40 chars"],
        "cta_buttons": ["3 CTA suggestions"]
    }},
    "instagram": {{
        "captions": [
            // 3 viral captions with emojis and hooks
        ],
        "hashtags": ["25 strategic hashtags"],
        "story_texts": ["3 story overlay texts"],
        "reels_hooks": ["5 viral reel hooks"]
    }},
    "seo": {{
        "titles": ["5 click-worthy SEO titles, 50-60 chars"],
        "meta_descriptions": ["5 compelling meta descriptions, 150-160 chars"],
        "keywords": {{
            "primary": ["5 primary keywords"],
            "secondary": ["10 secondary keywords"],
            "long_tail": ["10 long-tail phrases"]
        }}
    }},
    "landing_page": {{
        "hero_headline": "Powerful main headline",
        "hero_subheadline": "Supporting text",
        "value_props": ["4 benefit-focused value propositions"],
        "cta_texts": ["3 action-oriented CTAs"],
        "urgency_elements": ["3 urgency/scarcity elements"]
    }},
    "email": {{
        "subject_lines": [
            // 5 high-open-rate subject lines (use curiosity, numbers, urgency)
        ],
        "preview_texts": ["3 preview texts"],
        "cta_buttons": ["3 email CTAs"]
    }},
    "general": {{
        "taglines": ["5 memorable brand taglines"],
        "elevator_pitch": "30-second compelling pitch",
        "unique_selling_points": ["3 clear USPs"]
    }}
}}

Return ONLY valid JSON, no markdown or extra text.
"""

# =============================================================================
# LLM CONTENT GENERATION ENGINE (GROQ - FREE)
# =============================================================================

class ContentGenerator:
    """Main content generation engine using Groq (FREE & FAST)"""
    
    def __init__(self, api_key):
        self.client = Groq(api_key=api_key)
        self.model = "llama-3.3-70b-versatile"
    
    def generate_content(self, prompt, max_tokens=4000):
        """Generate content using Groq API"""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert marketing copywriter who creates HIGH-CONVERTING content. Always respond with valid JSON only. No markdown, no code blocks, no explanations - just pure JSON that can be parsed directly."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                max_tokens=max_tokens,
                temperature=0.8
            )
            
            content = response.choices[0].message.content.strip()
            
            if content.startswith('```'):
                content = re.sub(r'^```json?\n?', '', content)
                content = re.sub(r'\n?```$', '', content)
            
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                content = json_match.group()
            
            return json.loads(content)
            
        except json.JSONDecodeError as e:
            st.error(f"Error parsing response: {e}")
            return None
        except Exception as e:
            st.error(f"API Error: {e}")
            return None
    
    def generate_google_ads(self, inputs):
        prompt = PromptTemplates.google_ads_prompt(inputs)
        return self.generate_content(prompt)
    
    def generate_social_media(self, inputs):
        prompt = PromptTemplates.facebook_instagram_prompt(inputs)
        return self.generate_content(prompt)
    
    def generate_seo_content(self, inputs):
        prompt = PromptTemplates.seo_content_prompt(inputs)
        return self.generate_content(prompt)
    
    def generate_landing_page(self, inputs):
        prompt = PromptTemplates.landing_page_prompt(inputs)
        return self.generate_content(prompt)
    
    def generate_all_platforms(self, inputs):
        prompt = PromptTemplates.multi_platform_prompt(inputs)
        return self.generate_content(prompt, max_tokens=6000)

# =============================================================================
# EXPORT FUNCTIONS
# =============================================================================

def export_to_docx(content_data, inputs):
    """Export generated content to Word document"""
    doc = Document()
    
    title = doc.add_heading('AI Generated Marketing Content', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Business Information', level=1)
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_items = [
        ('Business Name', inputs.get('business_name', '')),
        ('Business Type', inputs.get('business_type', '')),
        ('Product/Service', inputs.get('product_service', '')),
        ('Target Audience', inputs.get('target_audience', '')),
        ('Offer', inputs.get('offer', '')),
        ('Tone', inputs.get('tone', ''))
    ]
    
    for i, (label, value) in enumerate(info_items):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    def add_list_section(title, items):
        if items:
            doc.add_heading(title, level=2)
            for item in items:
                doc.add_paragraph(f"• {item}", style='List Bullet')
            doc.add_paragraph()
    
    def add_dict_section(title, data, level=2):
        if data:
            doc.add_heading(title, level=level)
            if isinstance(data, dict):
                for key, value in data.items():
                    if isinstance(value, list):
                        doc.add_heading(key.replace('_', ' ').title(), level=level+1)
                        for item in value:
                            if isinstance(item, dict):
                                for k, v in item.items():
                                    doc.add_paragraph(f"{k}: {v}")
                            else:
                                doc.add_paragraph(f"• {item}", style='List Bullet')
                    elif isinstance(value, dict):
                        add_dict_section(key.replace('_', ' ').title(), value, level+1)
                    else:
                        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
            doc.add_paragraph()
    
    if isinstance(content_data, dict):
        for section_name, section_content in content_data.items():
            section_title = section_name.replace('_', ' ').title()
            
            if isinstance(section_content, list):
                add_list_section(section_title, section_content)
            elif isinstance(section_content, dict):
                add_dict_section(section_title, section_content)
            else:
                doc.add_heading(section_title, level=2)
                doc.add_paragraph(str(section_content))
    
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer

def export_to_pdf(content_data, inputs):
    """Export generated content to PDF"""
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=A4,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=20,
        spaceAfter=10,
        textColor=colors.HexColor('#667eea')
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceBefore=15,
        spaceAfter=8
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceBefore=5,
        spaceAfter=5
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceBefore=3,
        spaceAfter=3
    )
    
    story = []
    
    story.append(Paragraph("AI Generated Marketing Content", title_style))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("Business Information", heading_style))
    
    info_data = [
        ['Field', 'Value'],
        ['Business Name', inputs.get('business_name', '')],
        ['Business Type', inputs.get('business_type', '')],
        ['Product/Service', inputs.get('product_service', '')],
        ['Target Audience', inputs.get('target_audience', '')],
        ['Offer', inputs.get('offer', '')],
        ['Tone', inputs.get('tone', '')]
    ]
    
    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9ff')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#667eea'))
    ]))
    
    story.append(info_table)
    story.append(Spacer(1, 20))
    
    def add_content_to_story(data, level=0):
        if isinstance(data, dict):
            for key, value in data.items():
                title = key.replace('_', ' ').title()
                
                if level == 0:
                    story.append(Paragraph(title, heading_style))
                else:
                    story.append(Paragraph(title, subheading_style))
                
                if isinstance(value, list):
                    for item in value:
                        if isinstance(item, dict):
                            for k, v in item.items():
                                text = str(v).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(Paragraph(f"<b>{k}:</b> {text}", bullet_style))
                        else:
                            text = str(item).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            story.append(Paragraph(f"• {text}", bullet_style))
                elif isinstance(value, dict):
                    add_content_to_story(value, level + 1)
                else:
                    text = str(value).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(text, body_style))
                
                story.append(Spacer(1, 10))
        elif isinstance(data, list):
            for item in data:
                text = str(item).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(f"• {text}", bullet_style))
    
    if content_data:
        add_content_to_story(content_data)
    
    story.append(Spacer(1, 30))
    story.append(Paragraph(
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        ParagraphStyle('Footer', parent=styles['Normal'], alignment=1, fontSize=8)
    ))
    
    doc.build(story)
    pdf_buffer.seek(0)
    
    return pdf_buffer

# =============================================================================
# STREAMLIT UI COMPONENTS (PREMIUM)
# =============================================================================

def render_sidebar():
    """Render the premium sidebar"""
    with st.sidebar:
        st.markdown("""
        <div style="text-align: center; padding: 1rem 0;">
            <div style="font-size: 3rem; margin-bottom: 0.5rem;">🚀</div>
            <h2 style="color: white; margin: 0; font-family: 'Poppins', sans-serif;">AI Sales Agent</h2>
            <p style="color: rgba(255,255,255,0.7); font-size: 0.85rem; margin-top: 0.3rem;">Premium Edition • FREE</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        page = st.radio(
            "Navigation",
            ["🏠 Home", "✨ Generate Content", "📊 Dashboard", "⚙️ Settings"],
            label_visibility="collapsed"
        )
        
        st.markdown("---")
        
        api_key = st.text_input(
            "🔑 Groq API Key (FREE)",
            type="password",
            value=st.session_state.get('api_key', ''),
            help="Get FREE API key from console.groq.com"
        )
        
        if api_key:
            st.session_state['api_key'] = api_key
            st.success("✅ Connected!")
        else:
            st.warning("⚠️ Enter API Key")
            st.markdown("[🔗 Get FREE Key](https://console.groq.com)")
        
        st.markdown("---")
        
        st.markdown("""
        <div style="background: rgba(255,255,255,0.1); border-radius: 12px; padding: 1rem; margin-top: 1rem;">
            <h4 style="color: white; margin: 0 0 0.5rem 0; font-size: 0.9rem;">✨ Why Groq?</h4>
            <ul style="color: rgba(255,255,255,0.8); font-size: 0.8rem; margin: 0; padding-left: 1.2rem;">
                <li>100% FREE Forever</li>
                <li>10x Faster than GPT</li>
                <li>Llama 3.3 70B Model</li>
                <li>No Credit Card</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        return page

def render_input_form():
    """Render the premium input form"""
    st.markdown("""
    <div style="background: white; border-radius: 20px; padding: 2rem; box-shadow: 0 10px 40px rgba(0,0,0,0.08); margin-bottom: 2rem;">
        <h3 style="color: #1a1a2e; margin: 0 0 0.5rem 0; font-family: 'Poppins', sans-serif;">📝 Business Details</h3>
        <p style="color: #666; margin: 0;">Fill in your business information to generate high-converting content</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        business_name = st.text_input(
            "🏢 Business Name *",
            placeholder="e.g., TechFlow Solutions"
        )
        
        business_type = st.selectbox(
            "📁 Business Type *",
            [
                "E-commerce",
                "SaaS",
                "Local Service",
                "Consulting",
                "Healthcare",
                "Education",
                "Real Estate",
                "Finance",
                "Food & Restaurant",
                "Fitness & Wellness",
                "Travel & Tourism",
                "Manufacturing",
                "Agency",
                "Coaching",
                "Other"
            ]
        )
        
        product_service = st.text_area(
            "🎯 Product/Service Description *",
            placeholder="Describe your product or service in detail. What problem does it solve? What makes it unique?",
            height=120
        )
    
    with col2:
        target_audience = st.text_area(
            "👥 Target Audience *",
            placeholder="e.g., Small business owners aged 25-45, tech-savvy, looking to automate their workflow and scale their business...",
            height=100
        )
        
        offer = st.text_input(
            "🎁 Current Offer/Promotion",
            placeholder="e.g., 50% off for first 100 customers, Free trial, Limited time offer"
        )
        
        tone = st.selectbox(
            "🎨 Content Tone *",
            ["Professional", "Emotional", "Exciting", "Urgent", "Friendly", "Luxury"]
        )
    
    st.markdown("---")
    
    st.markdown("### 🎯 Select Platforms")
    
    platform = st.multiselect(
        "Choose platforms for content generation",
        [
            "All Platforms",
            "Google Ads",
            "Facebook",
            "Instagram",
            "SEO Content",
            "Landing Page"
        ],
        default=["All Platforms"],
        label_visibility="collapsed"
    )
    
    # Platform tags display
    if platform:
        tags_html = ""
        for p in platform:
            if p == "Google Ads":
                tags_html += '<span class="platform-tag tag-google">🎯 Google Ads</span>'
            elif p == "Facebook":
                tags_html += '<span class="platform-tag tag-facebook">📘 Facebook</span>'
            elif p == "Instagram":
                tags_html += '<span class="platform-tag tag-instagram">📸 Instagram</span>'
            elif p == "SEO Content":
                tags_html += '<span class="platform-tag tag-seo">🔍 SEO</span>'
            elif p == "Landing Page":
                tags_html += '<span class="platform-tag tag-landing">🏠 Landing Page</span>'
            elif p == "All Platforms":
                tags_html += '<span class="platform-tag" style="background: linear-gradient(135deg, #667eea, #764ba2); color: white;">🚀 All Platforms</span>'
        
        st.markdown(f'<div style="margin: 1rem 0;">{tags_html}</div>', unsafe_allow_html=True)
    
    return {
        'business_name': business_name,
        'business_type': business_type,
        'product_service': product_service,
        'target_audience': target_audience,
        'offer': offer,
        'tone': tone,
        'platform': platform
    }

def display_content_results(results, platform_type):
    """Display generated content with clean card layout - single headline focus"""
    if not results:
        st.warning("No content generated. Please try again.")
        return
    
    st.markdown("""
    <div style="background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%); padding: 1.5rem 2rem; border-radius: 16px; margin: 2rem 0; text-align: center;">
        <h2 style="color: white; margin: 0; font-family: 'Poppins', sans-serif;">⚡ Generated Content</h2>
    </div>
    """, unsafe_allow_html=True)
    
    # Helper function to check if key contains hashtags
    def is_hashtag_field(key):
        return 'hashtag' in key.lower()
    
    # Helper function to get first/best item from list
    def get_best_item(items):
        if isinstance(items, list) and len(items) > 0:
            return items[0]
        return items
    
    # Helper function to display hashtags as comma-separated
    def format_hashtags(hashtags_list):
        if isinstance(hashtags_list, list):
            return ", ".join(str(tag) for tag in hashtags_list)
        return str(hashtags_list)
    
    # Extract key content from results
    headline = ""
    description = ""
    cta = ""
    hashtags = ""
    keywords = []
    
    # Additional content storage
    extra_content = {}
    
    if isinstance(results, dict):
        for section_name, section_content in results.items():
            if isinstance(section_content, dict):
                for key, value in section_content.items():
                    key_lower = key.lower()
                    
                    # Get headline (first one only)
                    if 'headline' in key_lower and not headline:
                        if isinstance(value, list) and len(value) > 0:
                            headline = value[0]
                        elif isinstance(value, str):
                            headline = value
                    
                    # Get description (first one only)
                    elif 'description' in key_lower or 'primary_text' in key_lower:
                        if isinstance(value, list) and len(value) > 0:
                            if not description:
                                description = value[0]
                        elif isinstance(value, str) and not description:
                            description = value
                    
                    # Get CTA
                    elif 'cta' in key_lower or 'call' in key_lower:
                        if isinstance(value, list) and len(value) > 0:
                            if not cta:
                                cta = value[0]
                        elif isinstance(value, str) and not cta:
                            cta = value
                    
                    # Get hashtags
                    elif 'hashtag' in key_lower:
                        if isinstance(value, list):
                            hashtags = format_hashtags(value)
                        elif isinstance(value, str):
                            hashtags = value
                    
                    # Get keywords
                    elif 'keyword' in key_lower:
                        if isinstance(value, list):
                            keywords.extend(value)
                        elif isinstance(value, dict):
                            for k, v in value.items():
                                if isinstance(v, list):
                                    keywords.extend(v)
                    
                    # Store other content
                    else:
                        if section_name not in extra_content:
                            extra_content[section_name] = {}
                        extra_content[section_name][key] = value
            
            elif isinstance(section_content, list):
                section_lower = section_name.lower()
                if 'headline' in section_lower and not headline:
                    headline = section_content[0] if len(section_content) > 0 else ""
                elif 'hashtag' in section_lower:
                    hashtags = format_hashtags(section_content)
                elif 'keyword' in section_lower:
                    keywords.extend(section_content)
    
    # Display main content in clean cards
    
    # 1. HEADLINE - Highlighted
    if headline:
        st.markdown(f"""
        <div style="background: white; border-left: 4px solid #667eea; border-radius: 8px; padding: 1.2rem 1.5rem; margin: 1rem 0; box-shadow: 0 2px 10px rgba(0,0,0,0.08);">
            <p style="color: #333; font-size: 0.85rem; font-weight: 600; margin: 0 0 0.5rem 0; text-transform: uppercase; letter-spacing: 0.5px;">
                <span style="margin-right: 0.5rem;">📰</span> Headline:
            </p>
            <p style="color: #667eea; font-size: 1.3rem; font-weight: 700; margin: 0; line-height: 1.4;">
                {headline}
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # 2. DESCRIPTION
    if description:
        st.markdown(f"""
        <div style="background: white; border-left: 4px solid #667eea; border-radius: 8px; padding: 1.2rem 1.5rem; margin: 1rem 0; box-shadow: 0 2px 10px rgba(0,0,0,0.08);">
            <p style="color: #333; font-size: 0.85rem; font-weight: 600; margin: 0 0 0.5rem 0; text-transform: uppercase; letter-spacing: 0.5px;">
                <span style="margin-right: 0.5rem;">📝</span> Description:
            </p>
            <p style="color: #1a1a2e; font-size: 1.05rem; margin: 0; line-height: 1.6;">
                {description}
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # 3. CTA
    if cta:
        st.markdown(f"""
        <div style="background: white; border-left: 4px solid #667eea; border-radius: 8px; padding: 1.2rem 1.5rem; margin: 1rem 0; box-shadow: 0 2px 10px rgba(0,0,0,0.08);">
            <p style="color: #333; font-size: 0.85rem; font-weight: 600; margin: 0 0 0.5rem 0; text-transform: uppercase; letter-spacing: 0.5px;">
                <span style="margin-right: 0.5rem;">🎯</span> CTA:
            </p>
            <p style="color: #11998e; font-size: 1.1rem; font-weight: 600; margin: 0;">
                {cta} →
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # 4. HASHTAGS - Comma separated
    if hashtags:
        st.markdown(f"""
        <div style="background: white; border-left: 4px solid #667eea; border-radius: 8px; padding: 1.2rem 1.5rem; margin: 1rem 0; box-shadow: 0 2px 10px rgba(0,0,0,0.08);">
            <p style="color: #333; font-size: 0.85rem; font-weight: 600; margin: 0 0 0.5rem 0; text-transform: uppercase; letter-spacing: 0.5px;">
                <span style="margin-right: 0.5rem;">#️⃣</span> Hashtags:
            </p>
            <p style="color: #764ba2; font-size: 1rem; margin: 0; line-height: 1.8; word-wrap: break-word;">
                {hashtags}
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # 5. KEYWORDS - Comma separated
    if keywords:
        unique_keywords = list(dict.fromkeys(keywords))[:20]  # Remove duplicates, limit to 20
        keywords_str = ", ".join(unique_keywords)
        st.markdown(f"""
        <div style="background: white; border-left: 4px solid #667eea; border-radius: 8px; padding: 1.2rem 1.5rem; margin: 1rem 0; box-shadow: 0 2px 10px rgba(0,0,0,0.08);">
            <p style="color: #333; font-size: 0.85rem; font-weight: 600; margin: 0 0 0.5rem 0; text-transform: uppercase; letter-spacing: 0.5px;">
                <span style="margin-right: 0.5rem;">🔑</span> Keywords:
            </p>
            <p style="color: #1a1a2e; font-size: 1rem; margin: 0; line-height: 1.8;">
                {keywords_str}
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # Show "View All Content" expander for complete data
    st.markdown("---")
    
    with st.expander("📋 View All Generated Content", expanded=False):
        if isinstance(results, dict):
            for section_name, section_content in results.items():
                st.markdown(f"### {section_name.replace('_', ' ').title()}")
                
                if isinstance(section_content, dict):
                    for key, value in section_content.items():
                        st.markdown(f"**{key.replace('_', ' ').title()}:**")
                        if isinstance(value, list):
                            if is_hashtag_field(key):
                                st.write(", ".join(str(v) for v in value))
                            else:
                                for i, item in enumerate(value, 1):
                                    st.write(f"{i}. {item}")
                        elif isinstance(value, dict):
                            for k, v in value.items():
                                if isinstance(v, list):
                                    st.write(f"**{k}:** {', '.join(str(x) for x in v)}")
                                else:
                                    st.write(f"**{k}:** {v}")
                        else:
                            st.write(value)
                        st.write("")
                
                elif isinstance(section_content, list):
                    if is_hashtag_field(section_name):
                        st.write(", ".join(str(v) for v in section_content))
                    else:
                        for i, item in enumerate(section_content, 1):
                            st.write(f"{i}. {item}")
                
                st.markdown("---")

def render_generate_page():
    """Render the content generation page"""
    st.markdown("""
    <div class="premium-header">
        <h1>✨ Generate Marketing Content</h1>
        <p>Create high-converting content that drives sales and engagement</p>
        <div class="premium-badge">⚡ Powered by Groq AI • 100% FREE</div>
    </div>
    """, unsafe_allow_html=True)
    
    if 'api_key' not in st.session_state or not st.session_state['api_key']:
        st.markdown("""
        <div style="background: linear-gradient(135deg, #fff3cd 0%, #ffeeba 100%); padding: 2rem; border-radius: 16px; text-align: center; margin: 2rem 0;">
            <h3 style="color: #856404; margin: 0;">⚠️ API Key Required</h3>
            <p style="color: #856404; margin: 1rem 0;">Please enter your FREE Groq API key in the sidebar to continue.</p>
            <a href="https://console.groq.com" target="_blank" style="display: inline-block; background: #856404; color: white; padding: 0.8rem 2rem; border-radius: 10px; text-decoration: none; font-weight: 600;">🔑 Get FREE API Key →</a>
        </div>
        """, unsafe_allow_html=True)
        return
    
    inputs = render_input_form()
    
    required_fields = ['business_name', 'business_type', 'product_service', 'target_audience']
    is_valid = all(inputs.get(field) for field in required_fields)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        generate_btn = st.button(
            "🚀 Generate High-Converting Content",
            type="primary",
            use_container_width=True,
            disabled=not is_valid
        )
    
    if not is_valid and generate_btn:
        st.error("⚠️ Please fill in all required fields (*)")
        return
    
    if generate_btn:
        with st.spinner("🔄 Creating your high-converting content... This is fast!"):
            try:
                generator = ContentGenerator(st.session_state['api_key'])
                
                results = {}
                platforms = inputs['platform']
                
                if "All Platforms" in platforms:
                    results = generator.generate_all_platforms(inputs)
                else:
                    if "Google Ads" in platforms:
                        results['google_ads'] = generator.generate_google_ads(inputs)
                    if "Facebook" in platforms or "Instagram" in platforms:
                        social_results = generator.generate_social_media(inputs)
                        if social_results:
                            results.update(social_results)
                    if "SEO Content" in platforms:
                        results['seo'] = generator.generate_seo_content(inputs)
                    if "Landing Page" in platforms:
                        results['landing_page'] = generator.generate_landing_page(inputs)
                
                if results:
                    st.session_state['last_results'] = results
                    st.session_state['last_inputs'] = inputs
                    
                    text_for_nlp = f"{inputs['product_service']} {inputs['target_audience']} {inputs['offer']}"
                    nlp_keywords = extract_keywords_nlp(text_for_nlp)
                    st.session_state['nlp_keywords'] = nlp_keywords
                    
                    inputs_for_db = inputs.copy()
                    inputs_for_db['platform'] = ', '.join(inputs['platform']) if isinstance(inputs['platform'], list) else inputs['platform']
                    
                    flat_outputs = {
                        'headlines': json.dumps(results.get('google_ads', {}).get('headlines', [])),
                        'descriptions': json.dumps(results.get('google_ads', {}).get('descriptions', [])),
                        'hashtags': json.dumps(results.get('instagram', {}).get('hashtags', [])),
                        'keywords': json.dumps(nlp_keywords),
                        'cta': json.dumps(results.get('google_ads', {}).get('cta_suggestions', [])),
                        'seo_title': json.dumps(results.get('seo', {}).get('titles', [])),
                        'meta_description': json.dumps(results.get('seo', {}).get('meta_descriptions', [])),
                        'landing_page_content': json.dumps(results.get('landing_page', {}))
                    }
                    save_to_history(1, inputs_for_db, flat_outputs)
                    
            except Exception as e:
                st.error(f"Error generating content: {str(e)}")
                return
        
        if 'last_results' in st.session_state and st.session_state['last_results']:
            st.markdown("---")
            
            if 'nlp_keywords' in st.session_state:
                st.markdown("""
                <div style="background: white; border-radius: 20px; padding: 2rem; box-shadow: 0 10px 40px rgba(0,0,0,0.08); margin: 2rem 0;">
                    <h3 style="color: #1a1a2e; margin: 0 0 1rem 0;">🔑 Extracted Keywords & Hashtags</h3>
                """, unsafe_allow_html=True)
                
                col1, col2 = st.columns(2)
                with col1:
                    keywords = st.session_state['nlp_keywords']
                    st.markdown("**Keywords:**")
                    st.write(", ".join(keywords))
                with col2:
                    hashtags = generate_hashtags(keywords)
                    st.markdown("**Hashtags:**")
                    st.write(" ".join(hashtags))
                
                st.markdown("</div>", unsafe_allow_html=True)
            
            display_content_results(st.session_state['last_results'], inputs['platform'])
            
            st.markdown("---")
            st.markdown("""
            <div style="background: white; border-radius: 20px; padding: 2rem; box-shadow: 0 10px 40px rgba(0,0,0,0.08); margin: 2rem 0;">
                <h3 style="color: #1a1a2e; margin: 0 0 1rem 0;">📥 Export Your Content</h3>
            </div>
            """, unsafe_allow_html=True)
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                docx_buffer = export_to_docx(
                    st.session_state['last_results'],
                    st.session_state['last_inputs']
                )
                st.download_button(
                    label="📄 Download DOCX",
                    data=docx_buffer,
                    file_name=f"marketing_content_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col2:
                pdf_buffer = export_to_pdf(
                    st.session_state['last_results'],
                    st.session_state['last_inputs']
                )
                st.download_button(
                    label="📕 Download PDF",
                    data=pdf_buffer,
                    file_name=f"marketing_content_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            
            with col3:
                json_str = json.dumps(st.session_state['last_results'], indent=2)
                st.download_button(
                    label="📋 Download JSON",
                    data=json_str,
                    file_name=f"marketing_content_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    use_container_width=True
                )

def render_dashboard():
    """Render the dashboard with history"""
    st.markdown("""
    <div class="premium-header">
        <h1>📊 Content Dashboard</h1>
        <p>Track and manage your generated marketing content</p>
    </div>
    """, unsafe_allow_html=True)
    
    history = get_user_history(1, limit=20)
    
    if not history:
        st.info("📭 No content history yet. Generate some content to see it here!")
        return
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(f"""
        <div class="stats-card">
            <div class="stats-number">{len(history)}</div>
            <div class="stats-label">Total Generated</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        platforms = [h['platform'] for h in history if h['platform']]
        st.markdown(f"""
        <div class="stats-card" style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);">
            <div class="stats-number">{len(set(platforms))}</div>
            <div class="stats-label">Platforms Used</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        tones = [h['tone'] for h in history if h['tone']]
        most_common_tone = max(set(tones), key=tones.count) if tones else "N/A"
        st.markdown(f"""
        <div class="stats-card" style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);">
            <div class="stats-number" style="font-size: 1.5rem;">{most_common_tone}</div>
            <div class="stats-label">Most Used Tone</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        today_count = len([h for h in history if h['created_at'] and datetime.now().strftime('%Y-%m-%d') in h['created_at']])
        st.markdown(f"""
        <div class="stats-card" style="background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%);">
            <div class="stats-number">{today_count}</div>
            <div class="stats-label">Generated Today</div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    st.markdown("### 📜 Recent Content")
    
    for i, record in enumerate(history[:10]):
        with st.expander(f"📌 {record['business_name']} - {record['created_at'][:10] if record['created_at'] else 'N/A'}"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown(f"**Business Type:** {record['business_type']}")
                st.markdown(f"**Platform:** {record['platform']}")
                st.markdown(f"**Tone:** {record['tone']}")
            
            with col2:
                st.markdown(f"**Target Audience:** {record['target_audience']}")
                st.markdown(f"**Offer:** {record['offer']}")
            
            if record['full_response']:
                try:
                    content = json.loads(record['full_response'])
                    st.json(content)
                except:
                    st.text(record['full_response'])

def render_settings():
    """Render settings page"""
    st.markdown("""
    <div class="premium-header">
        <h1>⚙️ Settings</h1>
        <p>Configure your AI Sales Agent preferences</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("### 🔑 API Configuration")
    
    api_key = st.text_input(
        "Groq API Key (FREE)",
        type="password",
        value=st.session_state.get('api_key', ''),
        help="Get your FREE API key from console.groq.com"
    )
    
    if api_key:
        st.session_state['api_key'] = api_key
    
    st.markdown("[🔗 Get FREE Groq API Key](https://console.groq.com)")
    
    model = st.selectbox(
        "AI Model",
        ["llama-3.3-70b-versatile", "llama-3.1-8b-instant", "mixtral-8x7b-32768"],
        help="Select the Groq model to use (all FREE!)"
    )
    
    st.session_state['model'] = model
    
    st.markdown("---")
    st.markdown("### 🎨 Content Preferences")
    
    default_tone = st.selectbox(
        "Default Tone",
        ["Professional", "Emotional", "Exciting", "Urgent", "Friendly", "Luxury"],
        index=0
    )
    
    st.session_state['default_tone'] = default_tone
    
    st.markdown("---")
    st.markdown("### 📊 Data Management")
    
    if st.button("🗑️ Clear All History", type="secondary"):
        conn = sqlite3.connect('sales_content.db')
        cursor = conn.cursor()
        cursor.execute("DELETE FROM content_history")
        conn.commit()
        conn.close()
        st.success("✅ History cleared successfully!")
        st.rerun()

def render_home():
    """Render premium home page"""
    st.markdown("""
    <div class="premium-header">
        <h1>🚀 AI Sales Copy Agent</h1>
        <p>Generate High-Converting Marketing Content in Seconds</p>
        <div class="premium-badge">⚡ 100% FREE • Powered by Groq AI</div>
    </div>
    """, unsafe_allow_html=True)
    
    # Feature cards
    st.markdown("### ✨ What You Can Create")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">🎯</div>
            <div class="feature-title">Google Ads</div>
            <div class="feature-desc">
                High-converting headlines (30 chars), descriptions (90 chars), keywords targeting, and negative keywords.
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">📱</div>
            <div class="feature-title">Social Media</div>
            <div class="feature-desc">
                Scroll-stopping Facebook ads, viral Instagram captions, strategic hashtags, and engagement hooks.
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">🔍</div>
            <div class="feature-title">SEO Content</div>
            <div class="feature-desc">
                Click-worthy titles, compelling meta descriptions, keyword research, and schema suggestions.
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">🏠</div>
            <div class="feature-title">Landing Pages</div>
            <div class="feature-desc">
                Conversion-optimized hero sections, value propositions, CTAs, FAQs, and urgency elements.
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">📧</div>
            <div class="feature-title">Email Marketing</div>
            <div class="feature-desc">
                High-open-rate subject lines, preview texts, and compelling email CTAs that drive clicks.
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">📥</div>
            <div class="feature-title">Export Options</div>
            <div class="feature-desc">
                Download your content as PDF, Word documents, or JSON for easy integration.
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # How it works
    st.markdown("### 🎬 How It Works")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("""
        <div class="step-card">
            <div class="step-number">1</div>
            <div class="step-title">Enter Details</div>
            <p style="color: #666; font-size: 0.9rem;">Provide your business info, product, and target audience</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="step-card">
            <div class="step-number">2</div>
            <div class="step-title">Select Platforms</div>
            <p style="color: #666; font-size: 0.9rem;">Choose which platforms you want content for</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="step-card">
            <div class="step-number">3</div>
            <div class="step-title">Generate</div>
            <p style="color: #666; font-size: 0.9rem;">AI creates optimized content in seconds</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown("""
        <div class="step-card">
            <div class="step-number">4</div>
            <div class="step-title">Export & Use</div>
            <p style="color: #666; font-size: 0.9rem;">Download and use in your campaigns</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # CTA
    st.markdown("""
    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 3rem 2rem; border-radius: 24px; text-align: center; margin: 2rem 0;">
        <h2 style="color: white; margin: 0; font-family: 'Poppins', sans-serif;">Ready to Create High-Converting Content?</h2>
        <p style="color: rgba(255,255,255,0.9); margin: 1rem 0 2rem 0; font-size: 1.1rem;">Get your FREE Groq API key and start generating in seconds!</p>
        <div style="display: flex; justify-content: center; gap: 1rem; flex-wrap: wrap;">
            <a href="https://console.groq.com" target="_blank" style="display: inline-block; background: white; color: #667eea; padding: 1rem 2rem; border-radius: 12px; text-decoration: none; font-weight: 700; font-size: 1rem;">🔑 Get FREE API Key</a>
        </div>
    </div>
    """, unsafe_allow_html=True)

# =============================================================================
# MAIN APPLICATION
# =============================================================================

def main():
    """Main application entry point"""
    st.set_page_config(
        page_title="AI Sales Copy Agent | FREE",
        page_icon="🚀",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Inject premium CSS
    st.markdown(PREMIUM_CSS, unsafe_allow_html=True)
    
    # Initialize database
    init_database()
    
    # Initialize session state
    if 'api_key' not in st.session_state:
        st.session_state['api_key'] = os.getenv('GROQ_API_KEY', '')
    
    if 'last_results' not in st.session_state:
        st.session_state['last_results'] = None
    
    if 'last_inputs' not in st.session_state:
        st.session_state['last_inputs'] = None
    
    # Render sidebar and get current page
    page = render_sidebar()
    
    # Render appropriate page
    if page == "🏠 Home":
        render_home()
    elif page == "✨ Generate Content":
        render_generate_page()
    elif page == "📊 Dashboard":
        render_dashboard()
    elif page == "⚙️ Settings":
        render_settings()


if __name__ == "__main__":
    main()
